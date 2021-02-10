import requests
from datetime import datetime
from datetime import date
from docx import Document
from docx.shared import Pt
import re

# authorlist =

document = Document()
document.add_heading('MSSL Planetary Group Recent Publications', 0)
subtitle = document.add_paragraph("Generated - " + str(date.today()))

currentyear = date.today().year
years = range(currentyear, currentyear - 5, -1)

lib_token = '6PTOklOwoNVpIZlkzfX2imWZS1wuJyZ4uCZxijQR'  # Needs Updating for new people
numberofrequestrows = 5000
r = requests.get("https://api.adsabs.harvard.edu/v1/biblib/libraries/MLEF82wzRk60fgarKZPCow",
                 headers={"Authorization": "Bearer " + lib_token},
                 params={"rows": numberofrequestrows, "fl": "author_norm,title,pub_raw,doi,pubdate,property,year",
                         "sort": "first_author asc"})
paperlist = r.json()['solr']['response']['docs']

for year in years:
    separator = ', '
    document.add_heading(str(year), level=2)
    document.add_heading("Refereed", level=3)
    for paper in paperlist:
        if int(paper['year']) != year:
            continue
        else:
            if "REFEREED" in paper['property']:
                citationstring = u'{}, {}, {}, https://dx.doi.org/{}, {:%b %y}.'.format(
                    separator.join(paper['author_norm']),
                    paper['title'][0],
                    paper['pub_raw'].replace("<NUMPAGES>", "").replace("</NUMPAGES>", ""),
                    paper['doi'][0], datetime.strptime(paper['pubdate'][:-3], "%Y-%m"))
                # f.write(citationstring + "\n")
                paragraph = document.add_paragraph(citationstring)
                paragraph.paragraph_format.space_before = Pt(12)
                paragraph.paragraph_format.space_after = Pt(12)

    # f.write("\n Non-Refereed \n\n")
    document.add_heading("Non-Refereed", level=3)
    for paper in paperlist:
        if int(paper['year']) != year:
            continue
        else:
            if "NOT REFEREED" in paper['property']:
                conferences = ["epsc", "epsc-dps", "egu", "american geophysical union", "cospar"]
                if any(x in paper['pub_raw'].lower() for x in conferences):
                    continue
                citationstring = u'{}, {}, {}, {}.'.format(separator.join(paper['author_norm']), paper['title'][0],
                                                           paper['pub_raw'].replace("<NUMPAGES>", "").replace(
                                                               "</NUMPAGES>", ""),
                                                           paper['pubdate'])
                citationstring = re.sub('<A href="(.+)"> (.+)</A>', r'\2', citationstring)
                citationstring = citationstring.replace("&amp;", "&")

                paragraph = document.add_paragraph(citationstring)
                paragraph.paragraph_format.space_before = Pt(12)
                paragraph.paragraph_format.space_after = Pt(12)
    document.add_page_break()

document.save('MSSLRecentPublications.docx')
